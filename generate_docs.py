import os
import glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, start=100, end=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_type, val in [('top', top), ('bottom', bottom), ('start', start), ('end', end), ('left', start), ('right', end)]:
        node = OxmlElement(f'w:{m_type}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_toc(doc):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    p.add_run("Note: If the Table of Contents is empty, right-click here and select 'Update Field' to generate it.").italic = True
    
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    doc.add_page_break()

def apply_global_styles(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "B.L.A.S.T. OCR Engine Documentation | Technical Reference"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)

def add_heading(doc, text, level):
    if level == 1 and len(doc.paragraphs) > 3:
        doc.add_page_break()
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)
        h.paragraph_format.space_after = Pt(12)
        h.paragraph_format.space_before = Pt(24)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x43, 0x38, 0xca)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.space_before = Pt(16)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.space_before = Pt(12)

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_infobox(doc, target_audience, description):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_background(cell, 'F3F4F6')
    set_cell_margins(cell, 150, 150, 150, 150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"🎯 Target Audience: {target_audience}\n")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    run.font.name = 'Segoe UI'
    
    run_desc = p.add_run(description)
    run_desc.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    run_desc.font.name = 'Segoe UI'
    run_desc.italic = True
    
    doc.add_paragraph()

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_background(cell, '1E1E1E')
    set_cell_margins(cell, 150, 150, 150, 150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    
    # Simple syntax styling for code structure readability
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(248, 248, 242)
    
    doc.add_paragraph()

def add_file_content(doc, title, file_path, desc="Source file content extracted verbatim."):
    add_heading(doc, f"📄 {title}", level=2)
    add_paragraph(doc, desc, italic=True)
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        add_code_block(doc, content)
    else:
        p = doc.add_paragraph(f"⚠️ File not found: {file_path}")
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

def generate_doc():
    doc = Document()
    
    core_props = doc.core_properties
    core_props.author = "devhms"
    core_props.title = "B.L.A.S.T. OCR Engine Documentation"
    core_props.subject = "Technical Reference"
    core_props.comments = "Generated automatically with Aesthetic Upgrades"
    
    apply_global_styles(doc)
    
    # --- COVER PAGE ---
    for _ in range(6): doc.add_paragraph()
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title_p.add_run("⚡ B.L.A.S.T. OCR Engine ⚡")
    trun.font.name = 'Arial'
    trun.font.size = Pt(36)
    trun.bold = True
    trun.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle_p.add_run("Complete Technical Documentation & Source Reference")
    srun.font.name = 'Segoe UI'
    srun.font.size = Pt(16)
    srun.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    
    for _ in range(12): doc.add_paragraph()
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mrun = meta_p.add_run(f"Date Generated: {datetime.now().strftime('%B %d, %Y')}\nAuthor: devhms | Branch: main")
    mrun.font.name = 'Segoe UI'
    mrun.font.size = Pt(12)
    mrun.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    
    doc.add_page_break()
    
    # --- TABLE OF CONTENTS ---
    add_toc(doc)
    
    # --- SECTIONS ---
    add_heading(doc, "Section 1 — Executive Summary", 1)
    add_infobox(doc, "USER / BUSINESS", "High-level overview of the project goal and outcomes.")
    add_paragraph(doc, "The B.L.A.S.T. OCR Engine is a Deterministic OCR Automation tool designed for large-scale, automated extraction of text from scanned images, PDFs, and PPTX files. It implements a 3-layer A.N.T. architecture for stability, error recovery, and ultimate performance, prioritizing determinism over guesswork.")

    add_heading(doc, "Section 2 — Quickstart Guide", 1)
    add_infobox(doc, "USER", "Immediate commands to deploy and run the system.")
    add_paragraph(doc, "To quickly get started with the B.L.A.S.T. OCR engine, use the following commands from your terminal:")
    add_code_block(doc, "# Run CLI OCR Pipeline\npython run.py --source input.pdf --output out_dir\n\n# Launch Web Dashboard\nstreamlit run run_gui.py")
    add_file_content(doc, "README.md", "README.md", "Project Main Entry Documentation")
    add_file_content(doc, "run.py", "run.py", "CLI Application Entry Point")

    add_heading(doc, "Section 3 — Project Overview & Philosophy", 1)
    add_infobox(doc, "ALL", "The foundational principles guiding B.L.A.S.T. development.")
    add_paragraph(doc, "The project strictly adheres to the A.N.T. (Architect, Navigate, Tool) philosophy, providing a robust modular architecture to manage dependencies and system state. The Source of Truth for the design intent is maintained in gemini.md.")
    add_file_content(doc, "gemini.md", "gemini.md", "Project Master Plan & Source of Truth")
    add_file_content(doc, "ARCHITECTURE.md", "ARCHITECTURE.md", "A.N.T. Design Philosophy")

    add_heading(doc, "Section 4 — Architecture & Data Flow", 1)
    add_infobox(doc, "DEVELOPER", "Details the extraction flow and logic decisions.")
    add_file_content(doc, "extraction_flow.md", "architecture/extraction_flow.md", "Pipeline Routing Logic")

    add_heading(doc, "Section 5 — Environment & Setup", 1)
    add_infobox(doc, "DEVELOPER / DEVOPS", "Required configurations, dependencies, and health checks.")
    add_file_content(doc, "requirements.txt", "requirements.txt", "Python Dependencies")
    add_file_content(doc, "packages.txt", "packages.txt", "System / OS Dependencies")
    add_file_content(doc, ".env.example", ".env.example", "Environment Variables Template")
    add_file_content(doc, "dll_check.py", "dll_check.py", "Windows Environment Integrity Check")
    add_file_content(doc, "verify_foundation.py", "verify_foundation.py", "System Health Verification Script")

    add_heading(doc, "Section 6 — Core: BlastPipeline", 1)
    add_infobox(doc, "DEVELOPER", "The central orchestration engine (The Navigator).")
    add_file_content(doc, "blast_ocr/pipeline.py", "blast_ocr/pipeline.py", "Orchestrates batching, processing, and output generation.")
    add_file_content(doc, "blast_ocr/main.py", "blast_ocr/main.py", "API/CLI wrapper for backend interactions.")

    add_heading(doc, "Section 7 — Core: RobustOCRExtractor", 1)
    add_infobox(doc, "DEVELOPER", "The core worker and extraction logic (The Tool).")
    add_file_content(doc, "blast_ocr/core/extractor.py", "blast_ocr/core/extractor.py", "Implements CV2 preprocessing and EasyOCR execution.")
    add_file_content(doc, "blast_ocr/core/worker.py", "blast_ocr/core/worker.py", "Worker thread allocation for single page extraction.")

    add_heading(doc, "Section 8 — Caching System", 1)
    add_infobox(doc, "DEVELOPER", "High-performance hashing and disk caching.")
    add_file_content(doc, "blast_ocr/cache/manager.py", "blast_ocr/cache/manager.py", "Deduplication and processing bypass logic.")

    add_heading(doc, "Section 9 — Storage Layer", 1)
    add_infobox(doc, "DEVELOPER", "Relational persistence for job statuses.")
    add_file_content(doc, "blast_ocr/storage/database.py", "blast_ocr/storage/database.py", "SQLAlchemy ORM setup and ThreadLocal sessions.")

    add_heading(doc, "Section 10 — User Interface", 1)
    add_infobox(doc, "DEVELOPER", "Streamlit Dashboard components and premium styling.")
    add_file_content(doc, "blast_ocr/ui/web_app.py", "blast_ocr/ui/web_app.py", "Main Streamlit Dashboard Engine")
    add_file_content(doc, "blast_ocr/ui/styles.css", "blast_ocr/ui/styles.css", "Custom CSS Injection for Premium Glassmorphism Look")
    add_file_content(doc, ".streamlit/config.toml", ".streamlit/config.toml", "Streamlit Theming Engine")
    add_file_content(doc, "run_gui.py", "run_gui.py", "Streamlit Loader")

    add_heading(doc, "Section 11 — Self-Healing & Exceptions", 1)
    add_infobox(doc, "DEVELOPER", "Advanced error recovery strategies and custom exception taxonomies.")
    add_file_content(doc, "blast_ocr/core/healing.py", "blast_ocr/core/healing.py", "Retry loops and back-off architectures")
    add_file_content(doc, "blast_ocr/core/exceptions.py", "blast_ocr/core/exceptions.py", "Custom Exceptions and Error Classifications")

    add_heading(doc, "Section 12 — Parallel Processing", 1)
    add_infobox(doc, "DEVELOPER", "Multi-threading utilities tuned for optimal GPU/CPU load.")
    add_file_content(doc, "blast_ocr/core/parallel.py", "blast_ocr/core/parallel.py", "Manages multiple document workers securely")

    add_heading(doc, "Section 13 — Configuration Management", 1)
    add_infobox(doc, "DEVELOPER", "Type-safe environment configuration setup.")
    add_file_content(doc, "blast_ocr/config.py", "blast_ocr/config.py", "Pydantic Settings Models")
    add_file_content(doc, "blast_ocr/logging_config.py", "blast_ocr/logging_config.py", "Structured JSON and Desktop Loging Strategies")

    add_heading(doc, "Section 14 — Quality Assurance & Testing", 1)
    add_infobox(doc, "DEVELOPER / QA", "Unit tests ensuring core operations.")
    test_files = sorted(glob.glob("tests/*.py"))
    for tf in test_files:
        add_file_content(doc, os.path.basename(tf), tf, "QA Suite Component")

    add_heading(doc, "Section 15 — Maintenance & Audits", 1)
    add_infobox(doc, "DEVELOPER", "Project maintenance utilities to track system erosion.")
    add_file_content(doc, "AUDIT.md", "AUDIT.md", "Performance and Memory Leak Audit")
    add_file_content(doc, "maintain.py", "maintain.py", "Cleanup and health check automations")
    add_file_content(doc, "inventory_gen.py", "inventory_gen.py", "Workspace context gathering utility")

    add_heading(doc, "Section 16 — Enhancement Tracking", 1)
    add_infobox(doc, "DEVELOPER", "Records of continuous UX and Performance upgrades.")
    add_file_content(doc, "ENHANCEMENTS.md", "ENHANCEMENTS.md", "Visual and Logic upgrades")
    add_file_content(doc, "CHANGELOG.md", "CHANGELOG.md", "Historical adjustments")

    add_heading(doc, "Section 17 — Contribution Guidelines", 1)
    add_infobox(doc, "ALL", "Ground rules for introducing new features.")
    add_file_content(doc, "CONTRIBUTING.md", "CONTRIBUTING.md", "Rules of Engagement")
    add_file_content(doc, "DEVTOOLS_GUIDE.md", "DEVTOOLS_GUIDE.md", "Frontend Debugging Patterns")

    add_heading(doc, "Section 18 — Skill Knowledge Base", 1)
    add_infobox(doc, "DEVELOPER", "Distilled logic and patterns for maintaining the engine.")
    skill_files = sorted(glob.glob("skills/*.md"))
    for sf in skill_files:
        add_file_content(doc, os.path.basename(sf), sf, "System Context Meta-Skill")

    add_heading(doc, "Section 19 — Deployment & Packaging", 1)
    add_infobox(doc, "DEVOPS", "Containerization and shipment patterns.")
    add_paragraph(doc, "System dependencies are explicitly mapped in Section 5 (Environment). Ensure appropriate Docker/Poppler binaries are provisioned downstream.")

    add_heading(doc, "Section 20 — Benchmarking", 1)
    add_infobox(doc, "DEVELOPER", "Metrics verification scripts.")
    add_file_content(doc, "benchmark.py", "benchmark.py", "Measures execution speeds and limits.")

    add_heading(doc, "Section 21 — Appendix: Full Source Code Roster", 1)
    add_infobox(doc, "ALL", "Complete snapshot of the repository map.")
    
    tree_str = ""
    for root, dirs, files in os.walk("."):
        if ".git" in root or "__pycache__" in root or ".pytest_cache" in root or "output" in root: continue
        level = root.replace(".", "").count(os.sep)
        indent = " " * 4 * (level)
        tree_str += f"{indent}{os.path.basename(root)}/\n"
        subindent = " " * 4 * (level + 1)
        for f in files:
            tree_str += f"{subindent}{f}\n"

    add_code_block(doc, tree_str)

    output_dir = "C:/mnt/user-data/outputs"
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
        except Exception:
            output_dir = os.path.abspath("outputs")
            os.makedirs(output_dir, exist_ok=True)
    
    out_path = os.path.join(output_dir, "BLAST_OCR_Full_Documentation.docx")
    try:
        doc.save(out_path)
        print(f"SUCCESS: Premium Documentation successfully generated at {out_path}")
    except Exception as e:
        alt_path = "BLAST_OCR_Premium_Documentation.docx"
        doc.save(alt_path)
        print(f"FAILED to save to {out_path}, saved to {alt_path} instead. Error: {e}")

if __name__ == "__main__":
    generate_doc()
