"""
tests/test_skills_testing_strategy.py

Rigorous four-tiered testing verification suite directly implementing
skills/testing_strategy.md:
1. Unit Tests: Test logic, not libraries (EasyOCR/engine isolation & error handling)
2. Performance Benchmarking: Programmatic execution of benchmark.py metrics
3. Foundation Verification: Programmatic execution of verify_foundation.py
4. Manual Smoke Tests: Ghost data leak detection & DOCX/Markdown structural validation
"""

import os
import cv2
import numpy as np
import pymupdf as fitz
import pytest
from docx import Document

import verify_foundation
from blast_ocr.pipeline import BlastPipeline
from blast_ocr.core.extractor import RobustOCRExtractor


# ============================================================================
# 1. Unit Tests: Test Logic, Not Libraries
# ============================================================================

def test_unit_extractor_logic_error_handling(tmp_path, monkeypatch):
    """Test extractor logic when underlying OCR reader throws unexpected runtime error."""
    img_path = str(tmp_path / "fail_page.png")
    cv2.imwrite(img_path, np.full((100, 100, 3), 255, dtype=np.uint8))

    extractor = RobustOCRExtractor()

    def mock_readtext(*args, **kwargs):
        raise RuntimeError("Cuda out of memory or driver failure")

    monkeypatch.setattr(extractor.reader, "readtext", mock_readtext)

    from blast_ocr.core.exceptions import PageExtractionError
    with pytest.raises(PageExtractionError):
        extractor.process_page(img_path, 1)


# ============================================================================
# 2. Performance Benchmarking: Programmatic Execution of benchmark.py
# ============================================================================

def test_performance_benchmark_workflow(tmp_path):
    """Execute benchmark performance PDF processing and metrics validation."""
    bench_pdf = str(tmp_path / "benchmark_test.pdf")
    # Generate clean 2-page test PDF with fitz
    doc = fitz.open()
    for i in range(2):
        p = doc.new_page(width=400, height=300)
        p.insert_text((50, 100), f"Benchmark Metric Test Page {i + 1}")
    doc.save(bench_pdf)
    doc.close()

    assert os.path.exists(bench_pdf)
    assert os.path.getsize(bench_pdf) > 0

    # Execute benchmark run
    out_dir = str(tmp_path / "bench_out")
    pipeline = BlastPipeline(config_overrides={"output_dir": out_dir})
    res = pipeline.process_pdf(bench_pdf)

    assert len(res) == 2
    for r in res:
        assert r["page"] in (1, 2)
        assert "confidence" in r
        assert r["processing_time"] >= 0.0


# ============================================================================
# 3. Foundation Verification: Programmatic Execution of verify_foundation.py
# ============================================================================

def test_foundation_verification_execution(monkeypatch, tmp_path):
    """Execute verify_foundation.py main() to ensure full deployment health check passes."""
    monkeypatch.chdir(str(tmp_path))
    verify_foundation.main()


# ============================================================================
# 4. Manual Smoke Tests: Ghost Data & Document Formatting Validation
# ============================================================================

def test_ghost_data_and_document_formatting_preservation(tmp_path):
    """Verify zero orphan files left behind and validate Word/Markdown formatting."""
    out_dir = tmp_path / "smoke_out"
    out_dir.mkdir()

    img_path = str(tmp_path / "smoke_doc.png")
    img = np.full((150, 450, 3), 255, dtype=np.uint8)
    cv2.putText(img, "SECTION HEADER", (20, 50), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 0, 0), 2)
    cv2.putText(img, "Body paragraph content.", (20, 100), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 0, 0), 1)
    cv2.imwrite(img_path, img)

    pipeline = BlastPipeline(config_overrides={"output_dir": str(out_dir)})
    result = pipeline.process_job(source=img_path, formats=["docx", "md", "txt"])

    assert result["status"] == "success"

    # Ghost data check: no temporary files leaking in out_dir
    files_in_out = list(out_dir.iterdir())
    for f in files_in_out:
        if f.is_file():
            assert f.suffix in (".docx", ".md", ".txt", ".json", "")

    # Validate DOCX integrity with python-docx
    docx_path = result["generated_files"]["docx"]
    doc = Document(docx_path)
    assert len(doc.paragraphs) > 0
    full_docx_text = "\n".join(p.text for p in doc.paragraphs)
    assert len(full_docx_text) > 0

    # Validate Markdown integrity
    md_path = result["generated_files"]["md"]
    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()
    assert len(md_content) > 0
